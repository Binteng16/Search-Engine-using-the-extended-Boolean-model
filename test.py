from itertools import chain
import math
import sys
import zipfile
import PyPDF2
import os
import docx
from PyQt5 import QtCore, QtWidgets
from PyQt5.QtGui import *
from PyQt5.QtWidgets import *
from PyQt5.uic import loadUi


class MainWindows(QMainWindow):
    def __init__(self):
        super(MainWindows, self).__init__()
        loadUi('tampilan_baru.ui', self)
        self.loadFile.clicked.connect(self.prosesDokumen)
        self.pushButton.clicked.connect(self.compare_documents)

        # Initialize an empty custom dictionary
        self.custom_dict_path = "dictionary.docx"
        self.custom_dict = self.read_dictionary_from_docx(self.custom_dict_path)

        # Initialize an empty custom stopwords set
        self.custom_stopwords_path = "stopwordlist.docx"
        self.custom_stopwords = self.read_stopword_from_docx(self.custom_stopwords_path)

    def read_dictionary_from_docx(self, custom_dict_path):
        custom_dict_doc = docx.Document(custom_dict_path)
        custom_dict = [paragraph.text.strip() for paragraph in custom_dict_doc.paragraphs if paragraph.text.strip()]
        return custom_dict

    def read_stopword_from_docx(self, custom_stopwords_path):
        custom_stopwords_doc = docx.Document(custom_stopwords_path)
        custom_stopwords = [paragraph.text.strip() for paragraph in custom_stopwords_doc.paragraphs if paragraph.text.strip()]
        return custom_stopwords

    def read_PDF(self, pdf_path):
        text = ""
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)
            for page_num in range(num_pages):
                page = pdf_reader.pages[page_num]
                text += page.extract_text()
        return text

    def read_DOCX(self, docx_path):
        doc = docx.Document(docx_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text

    def read_TXT(self, txt_path):
        try:
            with open(txt_path, 'r', encoding='utf-8') as file:
                text = file.read()
            return text
        except UnicodeDecodeError:
            with open(txt_path, 'rb') as file:
                text = file.read()
            text = text.decode('latin-1')
            return text
    
    def hapus_imbuhan(self, kata):

        removed_prefix = ""
        removed_suffix = ""
        removed_infix = ""

        # Aturan awalan
        awalan = ["meng", "meny", "men", "mem", "me", "peng", "peny", "pen", "pem", "pe", "di", "ter", "ke", "se"]
        for a in awalan:
            if kata.startswith(a):
                removed_prefix = a
                kata = kata[len(a):]
                break  # Stop checking if a match is found

        # Aturan akhiran
        akhiran = ["kan", "an", "i", "lah", "kah", "tah", "pun", "nya"]
        for a in akhiran:
            if kata.endswith(a):
                removed_suffix = a
                kata = kata[:-len(a)]
                break  # Stop checking if a match is found

        # Aturan sisipan (infix)
        sisipan = ["el", "er", "el"]
        for s in sisipan:
            if s in kata:
                removed_infix = s
                kata = kata.replace(s, '')

        return kata, removed_prefix, removed_suffix, removed_infix
        
    def prosesDokumen(self, dokumen):
        file_dialog = QFileDialog()
        file_path, _ = file_dialog.getOpenFileName(
            self, 'Open File', '', 'Text Files (*.txt *.pdf *.docx);;All Files (*)')

        if file_path:
            if file_path.lower().endswith(('.pdf', '.docx', '.txt')):
                if file_path.lower().endswith('.pdf'):
                    dokumen = self.read_PDF(file_path)
                elif file_path.lower().endswith('.docx'):
                    dokumen = self.read_DOCX(file_path)
                elif file_path.lower().endswith('.txt'):
                    dokumen = self.read_TXT(file_path)

                self.textBrowser.setPlainText(dokumen)
                #case folding
                cf_dokumen = dokumen.lower()
                cf_dokumen = ''.join([char for char in cf_dokumen if char.isalnum() or char.isspace()])
                self.textBrowser_2.setPlainText(cf_dokumen)
                #tokenizing
                token_dokumen = cf_dokumen.split()
                self.textBrowser_3.setPlainText("\n".join(token_dokumen))
                #filtering
                filter_dokumen = [word for word in token_dokumen if word.lower() not in self.custom_stopwords]
                self.textBrowser_4.setPlainText("\n".join(filter_dokumen))
                #stemming
                removed_affix_words = [self.hapus_imbuhan(word) for word in filter_dokumen]
                self.textBrowser_5.setPlainText('\n'.join([f"{word[0]} ({word[1]}, {word[2]}, {word[3]})" for word in removed_affix_words]))
            else:
                self.showText.setPlainText("Unsupported file type.")
                
    def hitung_kata_dan_kuantitas(direktori):
        # Inisialisasi kamus untuk menyimpan jumlah kata
        jumlah_kata = {}

        # Iterasi melalui setiap dokumen dalam direktori
        for nama_file in os.listdir(direktori):
            path_file = os.path.join(direktori, nama_file)

            # Mengecek apakah path_file adalah file (bukan direktori)
            if os.path.isfile(path_file):
                # Membaca isi dokumen
                with open(path_file, 'r', encoding='utf-8') as file:
                    # Membaca kata-kata dari dokumen
                    kata_dokumen = file.read().split()

                    # Menghitung jumlah kata dalam dokumen
                    for kata in kata_dokumen:
                        # Mengupdate kamus jumlah kata
                        if kata in jumlah_kata:
                            jumlah_kata[kata] += 1
                        else:
                            jumlah_kata[kata] = 1

        # Mengembalikan hasil berupa kamus yang berisi kata dan jumlahnya
        return jumlah_kata

    # Contoh penggunaan
    direktori_dokumen = '/path/to/your/documents'
    hasil_perhitungan = hitung_kata_dan_kuantitas(direktori_dokumen)

    # Menampilkan hasil
    for kata, jumlah in hasil_perhitungan.items():
        print(f'{kata}: {jumlah}')
    # def compare_documents(self): 
    #     # Get ZIP file path
    #     zip_dialog = QFileDialog()
    #     zip_path, _ = zip_dialog.getOpenFileName(self, 'Open ZIP File', '', 'ZIP Files (*.zip)')

    #     if zip_path:
    #         try:
    #             with zipfile.ZipFile(zip_path, 'r') as zip_ref:
    #                 documents = []
    #                 for filename in zip_ref.namelist():
    #                     if filename.lower().endswith(('.pdf', '.docx', '.txt')):
    #                         with zip_ref.open(filename, 'r') as file:  # Read as text using 'r' mode
    #                             text = file.read().decode('utf-8')  # Decode bytes to text
    #                             documents.append(text)

    #                             # TF-IDF calculation (the rest of your TF-IDF code)

    #         except zipfile.BadZipFile:
    #             self.showText.setPlainText("Invalid ZIP file.")
    #         except Exception as e:
    #             self.showText.setPlainText(f"Error processing ZIP file: {e}")
    #     else:
    #         self.showText.setPlainText("No ZIP file selected.")
            
if __name__ == "__main__":
    app = QtWidgets.QApplication(sys.argv)
    window = MainWindows()
    window.setWindowTitle('Tugas Besar DataMining')
    window.show()
    sys.exit(app.exec_())
