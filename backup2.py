import sys
import os
import docx
import fitz
from PyQt5 import QtWidgets
from PyQt5.QtWidgets import QMainWindow, QFileDialog
from PyQt5.uic import loadUi

class MainWindows(QMainWindow):
    def __init__(self):
        super(MainWindows, self).__init__()
        loadUi('tampilan_baru.ui', self)
        self.loadFile.clicked.connect(self.prosesDokumen)
        self.pushButton.clicked.connect(self.showResult)

        # Initialize an empty custom dictionary
        self.custom_dict_path = "dictionary.docx"
        self.custom_dict = self.read_dictionary_from_docx(self.custom_dict_path)

        # Initialize an empty custom stopwords set
        self.custom_stopwords_path = "stopwordlist.docx"
        self.custom_stopwords = self.read_stopword_from_docx(self.custom_stopwords_path)

    def read_dictionary_from_docx(self, custom_dict_path):
        try:
            custom_dict_doc = docx.Document(custom_dict_path)
            custom_dict = [paragraph.text.strip() for paragraph in custom_dict_doc.paragraphs if paragraph.text.strip()]
            return custom_dict
        except Exception as e:
            print(f"Error reading dictionary: {e}")
            return []

    def read_stopword_from_docx(self, custom_stopwords_path):
        try:
            custom_stopwords_doc = docx.Document(custom_stopwords_path)
            custom_stopwords = [paragraph.text.strip() for paragraph in custom_stopwords_doc.paragraphs if paragraph.text.strip()]
            return custom_stopwords
        except Exception as e:
            print(f"Error reading stopwords: {e}")
            return []

    def read_text_from_docx(self, file_path):
        doc = docx.Document(file_path)
        return ' '.join([paragraph.text for paragraph in doc.paragraphs])

    def read_text_from_pdf(self, file_path):
        doc = fitz.open(file_path)
        text = ''
        for page_num in range(doc.page_count):
            page = doc[page_num]
            text += page.get_text()
        return text

    def read_text_from_txt(self, file_path):
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                return file.read()
        except Exception as e:
            print(f"Error reading text file: {e}")
            return ''

    def hapus_imbuhan(self, kata):
        removed_prefix = ""
        removed_suffix = ""
        removed_infix = ""

        # Aturan awalan
        awalan = ["meng", "meny", "men", "mem", "me", "peng", "peny", "pen", "pem", "pe", "di", "ter", "ke", "se"]
        for a in awalan:
            if kata.startswith(a):
                removed_prefix = a
                kata = kata[len(a):]
                break  # Stop checking if a match is found

        # Aturan akhiran
        akhiran = ["kan", "an", "i", "lah", "kah", "tah", "pun", "nya"]
        for a in akhiran:
            if kata.endswith(a):
                removed_suffix = a
                kata = kata[:-len(a)]
                break  # Stop checking if a match is found

        # Aturan sisipan (infix)
        sisipan = ["el", "er", "el"]
        for s in sisipan:
            if s in kata:
                removed_infix = s
                kata = kata.replace(s, '')

        return kata, removed_prefix, removed_suffix, removed_infix

    def prosesDokumen(self):
        file_dialog = QFileDialog()
        file_path, _ = file_dialog.getOpenFileName(
            self, 'Open File', '', 'Text Files (*.txt *.pdf *.docx);;All Files (*)')

        if file_path:
            if file_path.lower().endswith(('.pdf', '.docx', '.txt')):
                if file_path.lower().endswith('.pdf'):
                    dokumen = self.read_text_from_pdf(file_path)
                elif file_path.lower().endswith('.docx'):
                    dokumen = self.read_text_from_docx(file_path)
                elif file_path.lower().endswith('.txt'):
                    dokumen = self.read_text_from_txt(file_path)

                self.textBrowser.setPlainText(dokumen)
                # case folding
                cf_dokumen = dokumen.lower()
                cf_dokumen = ''.join([char for char in cf_dokumen if char.isalnum() or char.isspace()])
                self.textBrowser_2.setPlainText(cf_dokumen)
                # tokenizing
                token_dokumen = cf_dokumen.split()
                self.textBrowser_3.setPlainText("\n".join(token_dokumen))
                # filtering
                filter_dokumen = [word for word in token_dokumen if word.lower() not in self.custom_stopwords]
                self.textBrowser_4.setPlainText("\n".join(filter_dokumen))
                # stemming
                removed_affix_words = [self.hapus_imbuhan(word) for word in filter_dokumen]
                self.textBrowser_5.setPlainText(
                    '\n'.join([f"{word[0]} ({word[1]}, {word[2]}, {word[3]})" for word in removed_affix_words]))
            else:
                self.showText.setPlainText("Unsupported file type.")

    def nilaiDF(self, direktori): #df
        # Inisialisasi kamus untuk menyimpan jumlah kata dan dokumen
        jumlah_kata_dokumen = {}

        # Iterasi melalui setiap dokumen dalam direktori
        for nama_file in os.listdir(direktori):
            path_file = os.path.join(direktori, nama_file)

            # Mengecek apakah path_file adalah file (bukan direktori)
            if os.path.isfile(path_file):
                # Membaca isi dokumen berdasarkan jenis file
                if path_file.lower().endswith('.docx'):
                    kata_dokumen = self.read_text_from_docx(path_file).lower().split()
                elif path_file.lower().endswith('.pdf'):
                    kata_dokumen = self.read_text_from_pdf(path_file).lower().split()
                elif path_file.lower().endswith('.txt'):
                    kata_dokumen = self.read_text_from_txt(path_file).lower().split()
                else:
                    continue  # Skip unsupported file types

                # Menghapus karakter yang tidak diinginkan
                kata_dokumen = [''.join([char for char in kata if char.isalnum() or char.isspace()]) for kata in
                                kata_dokumen]
                # Menghapus stop words
                kata_dokumen = [kata for kata in kata_dokumen if kata.lower() not in self.custom_stopwords]
                # Menghapus imbuhan
                kata_dokumen = [self.hapus_imbuhan(kata) for kata in kata_dokumen]

                # Menghitung jumlah kata dalam dokumen
                for kata in set(kata_dokumen):  # Using set to count each word only once per document
                    # Mengupdate kamus jumlah kata dan dokumen
                    if kata in jumlah_kata_dokumen:
                        jumlah_kata_dokumen[kata][0] += 1  # Increment the word count
                        jumlah_kata_dokumen[kata][1].add(nama_file)  # Add the document to the set
                    else:
                        jumlah_kata_dokumen[kata] = [1, {nama_file}]  # Initialize with count 1 and document set

        # Mengembalikan hasil berupa kamus yang berisi kata, jumlahnya, dan jumlah dokumen yang memilikinya
        return jumlah_kata_dokumen
    
    def nilaiTF(self, direktori):
        # Inisialisasi kamus untuk menyimpan jumlah kata per dokumen
        term_frequency_per_dokumen = {}

        # Iterasi melalui setiap dokumen dalam direktori
        for nama_file in os.listdir(direktori):
            path_file = os.path.join(direktori, nama_file)

            # Mengecek apakah path_file adalah file (bukan direktori)
            if os.path.isfile(path_file):
                # Membaca isi dokumen berdasarkan jenis file
                if path_file.lower().endswith('.pdf'):
                    kata_dokumen = self.read_text_from_pdf(path_file).lower().split()
                elif path_file.lower().endswith('.docx'):
                    kata_dokumen = self.read_text_from_docx(path_file).lower().split()
                elif path_file.lower().endswith('.txt'):
                    kata_dokumen = self.read_text_from_txt(path_file).lower().split()
                else:
                    continue  # Skip unsupported file types

                # Menghapus karakter yang tidak diinginkan
                kata_dokumen = [''.join([char for char in kata if char.isalnum() or char.isspace()]) for kata in kata_dokumen]
                # Menghapus stop words
                kata_dokumen = [kata for kata in kata_dokumen if kata.lower() not in self.custom_stopwords]
                # Menghapus imbuhan
                kata_dokumen = [self.hapus_imbuhan(kata) for kata in kata_dokumen]

                # Menghitung jumlah kata dalam dokumen
                jumlah_kata = {}
                for kata in kata_dokumen:
                    # Mengupdate kamus jumlah kata
                    if kata in jumlah_kata:
                        jumlah_kata[kata] += 1
                    else:
                        jumlah_kata[kata] = 1

                # Menyimpan term frequency per dokumen
                term_frequency_per_dokumen[nama_file] = jumlah_kata

        # Mengembalikan hasil berupa kamus yang berisi term frequency per dokumen
        return term_frequency_per_dokumen
        
    def prosesQuery(self, query):
        # case folding
        cf_query = query.lower()
        cf_query = ''.join([char for char in cf_query if char.isalnum() or char.isspace()])
        # print(cf_query)
        # tokenizing
        token_query = cf_query.split()
        # print(token_query)
        # filtering
        filter_query = [word for word in token_query if word.lower() not in self.custom_stopwords]
        # print(filter_query)
        # stemming
        removed_affix_words = [self.hapus_imbuhan(word) for word in filter_query]
        # print(removed_affix_words)
        return removed_affix_words

    # def hitungTf(self, hasil_dokumen, hasil_query):
    #     # Combine hasil_dokumen dan hasil_query
    #     hasil_gabungan = dict(hasil_dokumen)
    #     for kata in hasil_query:
    #         hasil_gabungan[kata[0]] = hasil_gabungan.get(kata[0], 0) + 1

    #     jumlah_kata = sum(hasil_gabungan.values())
    #     tf = {kata: jumlah / jumlah_kata for kata, jumlah in hasil_gabungan.items()}
    #     return tf

    # def hitungdf(self, hasil_perhitungan):
    #     # Inisialisasi dictionary untuk menyimpan DF setiap kata
    #     df = {}

    #     # Set untuk melacak keberadaan kata pada dokumen
    #     kata_per_dokumen = set()

    #     # Iterasi melalui setiap dokumen
    #     for dokumen in hasil_perhitungan.values():
    #         # Set untuk melacak keberadaan kata pada dokumen tertentu
    #         kata_per_dokumen_dok = set()

    #         # Iterasi melalui setiap kata dalam dokumen
    #         for kata, _ in dokumen.items():
    #             # Jika kata belum ditambahkan ke kata_per_dokumen_dok, tambahkan nilai DF
    #             if kata not in kata_per_dokumen_dok:
    #                 df[kata] = df.get(kata, 0) + 1
    #                 kata_per_dokumen_dok.add(kata)

    #         # Gabungkan kata_per_dokumen_dok ke dalam kata_per_dokumen (untuk menghitung DF secara global)
    #         kata_per_dokumen.update(kata_per_dokumen_dok)

    #     # Mengembalikan nilai DF
    #     return df
    
    def showResult(self):
        # Contoh penggunaan
        direktori_dokumen = 'D:\\ITENAS\\KULIAH\\TUGAS ITENAS\\SEMESTER 5\\DATA MINING\\last\\data'
        hasil_tf = self.nilaiTF(direktori_dokumen)
        hasil_df = self.nilaiDF(direktori_dokumen)
        
        #query
        query = self.lineEdit.text()
        proses_query = self.prosesQuery(query)

        # Inisialisasi hasil_gabungan dengan hasil_df
        # gabung_hasil = dict(menghitung_kemunculan_kata, **{k: v for k, v in zip(proses_query, [1] * len(proses_query))})
        # gabung_hasil = dict(hasil_tf)
        hasil_gabungan = dict(hasil_df)
            
        # Menambahkan hasil query ke dalam hasil_gabungan
        for kata in proses_query:
            hasil_gabungan[kata[0]] = hasil_gabungan.get(kata[0], 0) + 1
            
        # for kata in proses_query:
        #     gabung_hasil[kata[0]] = gabung_hasil.get(kata[0], 0) + 1
        for nama_file, term_frequency in hasil_tf.items():
            print(f'\n{nama_file}:')
            for kata, jumlah in term_frequency.items():
                print(f'    {kata[0]}: {jumlah}')
            print('=======================================================================')
            
        print('\nHASIL QUERY')
        for kata in proses_query:
            print(f"{kata[0]}")
        print('=======================================================================')
            
        print('\nHASIL DF')
        for kata, (jumlah_kata, jumlah_dokumen) in hasil_df.items():
            print(f'\nKata "{kata[0]}": \nNilai DF: {jumlah_kata}, \nDokumen yang memunculkan: {jumlah_dokumen}')
        print('=======================================================================')
        # Menampilkan hasil gabungan
        # print('\nHASIL TF')
        # for kata, jumlah in hasil_tf.items():
        #     print(f"{kata}: {jumlah}")
        
        # for kata, jumlah in menghitung_kemunculan_kata.items():
        #     print(f"{kata[0]}: {jumlah}")
            
        # for kata in proses_query:
        #     hasil_gabungan[kata[0]] = hasil_gabungan.get(kata[0], 0) + 1
        
        # # Menampilkan hasil gabungan
        # for kata, jumlah in hasil_gabungan.items():
        #     print(f"{kata[0]}: {jumlah}")
            
        # # Menampilkan hasil gabungan
        # for kata, jumlah in hasil_gabungan.items():
        #     print(f"{kata[0]}: {jumlah}")
            
        # Menampilkan hasil TF   
        # tf = self.hitungTf(hasil_perhitungan_per_dokumen, proses_query)
        # for kata, nilai_tf in tf.items():
        #     print(f'TF value for "{kata}": {nilai_tf}')    
            
        # # Menampilkan hasil gabungan
        # for kata, jumlah in hasil_gabungan.items():
        #     print(f"{kata}: {jumlah}")

        # # Menampilkan hasil TF
        # tf = self.hitungTf(hasil_gabungan)
        # for kata, nilai_tf in tf.items():
        #     print(f'TF value for "{kata}": {nilai_tf}')

        # # Menampilkan hasil IDF
        # idf = self.hitungdf(hasil_perhitungan_per_dokumen)
        # for kata, nilai_idf in idf.items():
        #     print(f'IDF value for "{kata}": {nilai_idf}')


if __name__ == "__main__":
    app = QtWidgets.QApplication(sys.argv)
    window = MainWindows()
    window.setWindowTitle('Tugas Besar DataMining')
    window.show()
    sys.exit(app.exec_())
