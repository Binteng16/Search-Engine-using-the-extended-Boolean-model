import sys
import os
import docx
import fitz
import math
from PyQt5 import QtWidgets
from PyQt5.QtWidgets import QMainWindow, QFileDialog
from PyQt5.uic import loadUi

class MainWindows(QMainWindow):
    def __init__(self):
        super(MainWindows, self).__init__()
        loadUi('tampilan_baru.ui', self)
        self.loadFile.clicked.connect(self.prosesDokumen)
        self.pushButton.clicked.connect(self.showResult)

        # Initialize an empty custom dictionary
        self.custom_dict_path = "dictionary.docx"
        self.custom_dict = self.read_dictionary_from_docx(self.custom_dict_path)

        # Initialize an empty custom stopwords set
        self.custom_stopwords_path = "stopwordlist.docx"
        self.custom_stopwords = self.read_stopword_from_docx(self.custom_stopwords_path)

    def read_dictionary_from_docx(self, custom_dict_path):
        try:
            custom_dict_doc = docx.Document(custom_dict_path)
            custom_dict = [paragraph.text.strip() for paragraph in custom_dict_doc.paragraphs if paragraph.text.strip()]
            return custom_dict
        except Exception as e:
            print(f"Error reading dictionary: {e}")
            return []

    def read_stopword_from_docx(self, custom_stopwords_path):
        try:
            custom_stopwords_doc = docx.Document(custom_stopwords_path)
            custom_stopwords = [paragraph.text.strip() for paragraph in custom_stopwords_doc.paragraphs if paragraph.text.strip()]
            return custom_stopwords
        except Exception as e:
            print(f"Error reading stopwords: {e}")
            return []

    def read_text_from_docx(self, file_path):
        doc = docx.Document(file_path)
        return ' '.join([paragraph.text for paragraph in doc.paragraphs])

    def read_text_from_pdf(self, file_path):
        doc = fitz.open(file_path)
        text = ''
        for page_num in range(doc.page_count):
            page = doc[page_num]
            text += page.get_text()
        return text

    def read_text_from_txt(self, file_path):
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                return file.read()
        except Exception as e:
            print(f"Error reading text file: {e}")
            return ''

    def hapus_imbuhan(self, kata):
        removed_prefix = ""
        removed_suffix = ""
        removed_infix = ""

        # Aturan awalan
        awalan = ["meng", "meny", "men", "mem", "me", "peng", "peny", "pen", "pem", "pe", "di", "ter", "ke", "se"]
        for a in awalan:
            if kata.startswith(a):
                removed_prefix = a
                kata = kata[len(a):]
                break  # Stop checking if a match is found

        # Aturan akhiran
        akhiran = ["kan", "an", "i", "lah", "kah", "tah", "pun", "nya"]
        for a in akhiran:
            if kata.endswith(a):
                removed_suffix = a
                kata = kata[:-len(a)]
                break  # Stop checking if a match is found

        # Aturan sisipan (infix)
        sisipan = ["el", "er", "el"]
        for s in sisipan:
            if s in kata:
                removed_infix = s
                kata = kata.replace(s, '')

        return kata, removed_prefix, removed_suffix, removed_infix

    def prosesDokumen(self):
        file_dialog = QFileDialog()
        file_path, _ = file_dialog.getOpenFileName(
            self, 'Open File', '', 'Text Files (*.txt *.pdf *.docx);;All Files (*)')

        if file_path:
            if file_path.lower().endswith(('.pdf', '.docx', '.txt')):
                if file_path.lower().endswith('.pdf'):
                    dokumen = self.read_text_from_pdf(file_path)
                elif file_path.lower().endswith('.docx'):
                    dokumen = self.read_text_from_docx(file_path)
                elif file_path.lower().endswith('.txt'):
                    dokumen = self.read_text_from_txt(file_path)

                self.textBrowser.setPlainText(dokumen)
                # case folding
                cf_dokumen = dokumen.lower()
                cf_dokumen = ''.join([char for char in cf_dokumen if char.isalnum() or char.isspace()])
                self.textBrowser_2.setPlainText(cf_dokumen)
                # tokenizing
                token_dokumen = cf_dokumen.split()
                self.textBrowser_3.setPlainText("\n".join(token_dokumen))
                # filtering
                filter_dokumen = [word for word in token_dokumen if word.lower() not in self.custom_stopwords]
                self.textBrowser_4.setPlainText("\n".join(filter_dokumen))
                # stemming
                removed_affix_words = [self.hapus_imbuhan(word) for word in filter_dokumen]
                self.textBrowser_5.setPlainText(
                    '\n'.join([f"{word[0]} ({word[1]}, {word[2]}, {word[3]})" for word in removed_affix_words]))
            else:
                self.showText.setPlainText("Unsupported file type.")

    def nilaiDF(self, direktori):
        # Inisialisasi kamus untuk menyimpan jumlah kata dan dokumen
        jumlah_kata_dokumen = {}
        
        # Set untuk melacak nama dokumen yang telah dibaca
        dokumen_yang_dibaca = set()

        # Iterasi melalui setiap dokumen dalam direktori
        for nama_file in os.listdir(direktori):
            path_file = os.path.join(direktori, nama_file)

            # Mengecek apakah path_file adalah file (bukan direktori)
            if os.path.isfile(path_file):
                # Membaca isi dokumen berdasarkan jenis file
                if path_file.lower().endswith('.docx'):
                    kata_dokumen = self.read_text_from_docx(path_file).lower().split()
                elif path_file.lower().endswith('.pdf'):
                    kata_dokumen = self.read_text_from_pdf(path_file).lower().split()
                elif path_file.lower().endswith('.txt'):
                    kata_dokumen = self.read_text_from_txt(path_file).lower().split()
                else:
                    continue  # Skip unsupported file types

                # Menghapus karakter yang tidak diinginkan
                kata_dokumen = [''.join([char for char in kata if char.isalnum() or char.isspace()]) for kata in
                                kata_dokumen]
                # Menghapus stop words
                kata_dokumen = [kata for kata in kata_dokumen if kata.lower() not in self.custom_stopwords]
                # Menghapus imbuhan
                kata_dokumen = [self.hapus_imbuhan(kata) for kata in kata_dokumen]

                # Menghitung jumlah kata dalam dokumen
                for kata in set(kata_dokumen):  # Using set to count each word only once per document
                    # Mengupdate kamus jumlah kata dan dokumen
                    if kata in jumlah_kata_dokumen:
                        jumlah_kata_dokumen[kata][0] += 1  # Increment the word count
                        jumlah_kata_dokumen[kata][1].add(nama_file)  # Add the document to the set
                    else:
                        jumlah_kata_dokumen[kata] = [1, {nama_file}]  # Initialize with count 1 and document set

                # Menambahkan nama dokumen ke dalam set
                dokumen_yang_dibaca.add(nama_file)

        # Menghitung jumlah dokumen yang dibaca
        jumlah_dokumen_yang_dibaca = len(dokumen_yang_dibaca)

        # Menambahkan informasi jumlah dokumen yang dibaca ke dalam kamus
        jumlah_kata_dokumen["jumlah_dokumen_yang_dibaca"] = jumlah_dokumen_yang_dibaca

        # Mengembalikan hasil berupa kamus yang berisi kata, jumlahnya, dan jumlah dokumen yang memilikinya
        return jumlah_kata_dokumen

    
    def nilaiTF(self, direktori):
        # Inisialisasi kamus untuk menyimpan jumlah kata per dokumen
        term_frequency_per_dokumen = {}

        # Iterasi melalui setiap dokumen dalam direktori
        for nama_file in os.listdir(direktori):
            path_file = os.path.join(direktori, nama_file)

            # Mengecek apakah path_file adalah file (bukan direktori)
            if os.path.isfile(path_file):
                # Membaca isi dokumen berdasarkan jenis file
                if path_file.lower().endswith('.pdf'):
                    kata_dokumen = self.read_text_from_pdf(path_file).lower().split()
                elif path_file.lower().endswith('.docx'):
                    kata_dokumen = self.read_text_from_docx(path_file).lower().split()
                elif path_file.lower().endswith('.txt'):
                    kata_dokumen = self.read_text_from_txt(path_file).lower().split()
                else:
                    continue  # Skip unsupported file types

                # Menghapus karakter yang tidak diinginkan
                kata_dokumen = [''.join([char for char in kata if char.isalnum() or char.isspace()]) for kata in kata_dokumen]
                # Menghapus stop words
                kata_dokumen = [kata for kata in kata_dokumen if kata.lower() not in self.custom_stopwords]
                # Menghapus imbuhan
                kata_dokumen = [self.hapus_imbuhan(kata) for kata in kata_dokumen]

                # Menghitung jumlah kata dalam dokumen
                jumlah_kata = {}
                for kata in kata_dokumen:
                    # Mengupdate kamus jumlah kata
                    if kata in jumlah_kata:
                        jumlah_kata[kata] += 1
                    else:
                        jumlah_kata[kata] = 1

                # Menyimpan term frequency per dokumen
                term_frequency_per_dokumen[nama_file] = jumlah_kata

        # Mengembalikan hasil berupa kamus yang berisi term frequency per dokumen
        return term_frequency_per_dokumen
    
    def nilaiMaxTF(self, direktori):
        # Memanggil fungsi nilaiTF untuk mendapatkan term frequency per dokumen
        term_frequency_per_dokumen = self.nilaiTF(direktori)

        # Inisialisasi variabel untuk menyimpan nilai maksimum
        nilai_max_per_dokumen = {}

        # Iterasi melalui hasil term frequency per dokumen
        for nama_file, term_frequency in term_frequency_per_dokumen.items():
            # Mencari nilai maksimum jumlah kata dalam dokumen
            max_jumlah_kata = max(term_frequency.values())

            # Menyimpan nilai maksimum per dokumen
            nilai_max_per_dokumen[nama_file] = max_jumlah_kata

        # Mengembalikan hasil berupa kamus yang berisi nilai maksimum per dokumen
        return nilai_max_per_dokumen
        
    def prosesQuery(self, query):
        # case folding
        cf_query = query.lower()
        cf_query = ''.join([char for char in cf_query if char.isalnum() or char.isspace()])
        # print(cf_query)
        # tokenizing
        token_query = cf_query.split()
        # print(token_query)
        # filtering
        filter_query = [word for word in token_query if word.lower() not in self.custom_stopwords]
        # print(filter_query)
        # stemming
        removed_affix_words = [self.hapus_imbuhan(word) for word in filter_query]
        # print(removed_affix_words)
        return removed_affix_words
    
    def showResult(self):
        try:
            # Contoh penggunaan
            direktori_dokumen = 'D:\\ITENAS\\KULIAH\\TUGAS ITENAS\\SEMESTER 5\\DATA MINING\\last\\data'
            hasil_tf = self.nilaiTF(direktori_dokumen)
            hasil_df = self.nilaiDF(direktori_dokumen)
            hasil_max_tf = self.nilaiMaxTF(direktori_dokumen)
            
            #query
            query = self.lineEdit.text()
            proses_query = self.prosesQuery(query)
            hasil_gabungan = dict(hasil_df)
                
            # Menambahkan hasil query ke dalam hasil_gabungan
            for kata in proses_query:
                hasil_gabungan[kata[0]] = hasil_gabungan.get(kata[0], 0) + 1
            
            # Mengambil nilai jumlah_dokumen_yang_dibaca dari hasil nilaiDF
            jumlah_dokumen_yang_dibaca = hasil_df.get("jumlah_dokumen_yang_dibaca", 0)

            # Mencetak hanya baris yang berisi informasi jumlah dokumen yang dibaca
            print('\nHASIL QUERY')
            proses_query = self.prosesQuery(query)
            for kata in proses_query:
                print(f"{kata[0]}")

            # Menampilkan jumlah kata dari query
            total_kata_query = len(proses_query)
            print(f'\nJUMLAH KATA QUERY: {total_kata_query}')
            print('========================================= OR =========================================')
            print('\nQUERY DENGAN DOKUMEN (OR)')
            for nama_file, nilai_max in hasil_max_tf.items():
                print(f'\n{nama_file}:')
                
                # Pernyataan cetak untuk memeriksa kata-kata yang sesuai dengan query
                kata_query_sesuai = [kata[0] for kata in set(proses_query) & set(hasil_tf[nama_file].keys())]
                print(f'Kata-kata query yang sesuai dengan dokumen: {set(kata_query_sesuai)}')
                
                total_pangkat = 0
                for kata, jumlah in hasil_tf[nama_file].items():
                    kata_df, nilai_df, dokumen_set = kata[0], hasil_df[kata][0], hasil_df[kata][1]

                    # Mengambil hanya bagian kata yang sesuai dengan query
                    if kata_df in kata_query_sesuai:
                        hasil_pembagian_ntf = jumlah / nilai_max
                        hasil_perhitungan_nidf = math.log10(jumlah_dokumen_yang_dibaca / nilai_df) / math.log10(jumlah_dokumen_yang_dibaca)
                        hasil_ntf_nidf = hasil_pembagian_ntf * hasil_perhitungan_nidf
                        hasil_ntf_nidf_pangkatP = hasil_ntf_nidf ** 2
                        total_pangkat += hasil_ntf_nidf_pangkatP
                        print(f'\n{kata_df}:       \nPERHITUNGAN : {hasil_pembagian_ntf} * {hasil_perhitungan_nidf} = {hasil_ntf_nidf} ')
                        print(f'HASIL PANGKAT : {hasil_perhitungan_nidf} ^ 2 = {hasil_ntf_nidf_pangkatP}')
                print(f'TOTAL HASIL PANGKAT: {total_pangkat}')

                # Menampilkan hasil akhir (total_pangkat / total_kata_query)
                if total_kata_query > 0:
                    hasil_pembagian_1 = total_pangkat / total_kata_query
                    print(f'\nHASIL PEMBAGIAN: {total_pangkat} / {total_kata_query} = {hasil_pembagian_1}')
                else:
                    print('\nTidak ada kata dalam query.')
                hasil_akhir = math.sqrt(hasil_pembagian_1)
                print(f'HASIL AKHIR : {hasil_akhir}')
            print('========================================= AND =========================================')
            print('\nQUERY DENGAN DOKUMEN (AND)')
            for nama_file, nilai_max in hasil_max_tf.items():
                print(f'\n{nama_file}:')
                
                # Pernyataan cetak untuk memeriksa kata-kata yang sesuai dengan query
                kata_query_sesuai = [kata[0] for kata in set(proses_query) & set(hasil_tf[nama_file].keys())]
                print(f'Kata-kata query yang sesuai dengan dokumen: {set(kata_query_sesuai)}')
                
                total_pangkat = 0
                for kata, jumlah in hasil_tf[nama_file].items():
                    kata_df, nilai_df, dokumen_set = kata[0], hasil_df[kata][0], hasil_df[kata][1]

                    # Mengambil hanya bagian kata yang sesuai dengan query
                    if kata_df in kata_query_sesuai:
                        hasil_pembagian_ntf = jumlah / nilai_max
                        hasil_perhitungan_nidf = math.log10(jumlah_dokumen_yang_dibaca / nilai_df) / math.log10(jumlah_dokumen_yang_dibaca)
                        hasil_ntf_nidf = hasil_pembagian_ntf * hasil_perhitungan_nidf
                        hasil_ntf_nidf_pangkatP = (1 - hasil_ntf_nidf) ** 2
                        total_pangkat += hasil_ntf_nidf_pangkatP
                        print(f'\n{kata_df}:       \nPERHITUNGAN : {hasil_pembagian_ntf} * {hasil_perhitungan_nidf} = {hasil_ntf_nidf} ')
                        print(f'HASIL PANGKAT : (1 - {hasil_perhitungan_nidf}) ^ 2 = {hasil_ntf_nidf_pangkatP}')
                print(f'TOTAL HASIL PANGKAT: {total_pangkat}')

                # Menampilkan hasil akhir (total_pangkat / total_kata_query)
                if total_kata_query > 0:
                    hasil_pembagian_1 = total_pangkat / total_kata_query
                    print(f'\nHASIL PEMBAGIAN: {total_pangkat} / {total_kata_query} = {hasil_pembagian_1}')
                else:
                    print('\nTidak ada kata dalam query.')
                hasil_akhir = 1 - (math.sqrt(hasil_pembagian_1))
                print(f'HASIL AKHIR : {hasil_akhir}')    
            print('========================================= PENJELASAN =========================================')
            print('\nHASIL NTF x NIDF')
            for nama_file, nilai_max in hasil_max_tf.items():
                print(f'\n{nama_file}:')
                for kata, jumlah in hasil_tf[nama_file].items():
                    hasil_pembagian_ntf = jumlah / nilai_max
                    kata_df, nilai_df, dokumen_set = kata[0], hasil_df[kata][0], hasil_df[kata][1]
                    hasil_perhitungan_nidf = math.log10(jumlah_dokumen_yang_dibaca / nilai_df) / math.log10(jumlah_dokumen_yang_dibaca)
                    hasil_ntf_nidf = hasil_pembagian_ntf * hasil_perhitungan_nidf
                    print(f'{kata_df}:       \nPERHITUNGAN : {hasil_pembagian_ntf} * {hasil_perhitungan_nidf} = {hasil_ntf_nidf} ')
            print('=======================================================================')
            print('\nJUMLAH DOKUMEN YANG DIBACA')
            for kata, info_kata in hasil_df.items():
                if kata == "jumlah_dokumen_yang_dibaca":
                    jumlah_dokumen_yang_dibaca = info_kata
                    print(f'Dokumen yang Dibaca: {jumlah_dokumen_yang_dibaca}')
            print('=======================================================================')
            # Menampilkan hasil dari jumlah dokumen yang dibaca dibagi dengan nilai DF
            if jumlah_dokumen_yang_dibaca > 0:
                print('\nHASIL NIDF (JUMLAH DOKUMEN YANG DIBACA DIBAGI DENGAN NILAI DF)')
                for kata, info_kata in hasil_df.items():
                    if kata != "jumlah_dokumen_yang_dibaca":
                        if isinstance(info_kata, list):
                            kata_df, nilai_df, dokumen_set = kata[0], info_kata[0], info_kata[1]
                            hasil_perhitungan = math.log10(jumlah_dokumen_yang_dibaca / nilai_df) / math.log10(jumlah_dokumen_yang_dibaca)
                            print(f'\nKata "{kata_df}": \nNilai DF: {nilai_df}, \nDokumen yang memunculkan: {dokumen_set}, \nHasil Perhitungan: log[{jumlah_dokumen_yang_dibaca} / {nilai_df}] / log({jumlah_dokumen_yang_dibaca}) = {hasil_perhitungan}')
            else:
                print('\nTidak dapat melakukan perhitungan karena jumlah dokumen yang dibaca adalah 0.')
        
            print('=======================================================================')
            print('\nHASIL NTF (JUMLAH KATA DIBAGI DENGAN NILAI MAKSIMUM PER DOKUMEN)')
            for nama_file, nilai_max in hasil_max_tf.items():
                print(f'\n{nama_file}:')
                for kata, jumlah in hasil_tf[nama_file].items():
                    hasil_pembagian =  jumlah / nilai_max
                    print(f'    {kata[0]}: {jumlah} / {nilai_max} = {hasil_pembagian}')
            print('=======================================================================')
            print('\nNILAI MAKSIMUM TF')
            for nama_file, nilai_max in hasil_max_tf.items():
                print(f'{nama_file}: {nilai_max}')
            print('=======================================================================')
            print('\nHASIL TF')  
            for nama_file, term_frequency in hasil_tf.items():
                print(f'\n{nama_file}:')
                for kata, jumlah in term_frequency.items():
                    print(f'    {kata[0]}: {jumlah}')
            print('=======================================================================')
            print('\nHASIL DF')
            for kata, info_kata in hasil_df.items():
                # Menyaring kata "jumlah_dokumen_yang_dibaca"
                if kata != "jumlah_dokumen_yang_dibaca":
                    if isinstance(info_kata, int):
                        print(f'\nKata "{kata}": \nNilai DF: {info_kata}')
                    else:
                        print(f'\nKata "{kata[0]}": \nNilai DF: {info_kata[0]}, \nDokumen yang memunculkan: {info_kata[1]}')
            print('=======================================================================')
            

            
                
        except Exception as e:
            print(f"Error in showResult: {e}")

if __name__ == "__main__":
    app = QtWidgets.QApplication(sys.argv)
    window = MainWindows()
    window.setWindowTitle('Tugas Besar DataMining')
    window.show()
    sys.exit(app.exec_())